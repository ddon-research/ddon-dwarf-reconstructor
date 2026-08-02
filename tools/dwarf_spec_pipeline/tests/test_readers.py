from __future__ import annotations

from pathlib import Path

import pytest
from docx import Document

from dwarf_spec_pipeline.readers import read_docx, read_html


@pytest.mark.unit
def test_docx_reader_preserves_body_order_and_merged_cell_span(tmp_path: Path) -> None:
    document = Document()
    document.add_heading("1. Introduction", level=1)
    document.add_paragraph("Text before the table.")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).merge(table.cell(0, 1)).text = "Merged heading"
    table.cell(1, 0).text = "Name"
    table.cell(1, 1).text = "Value"
    document.add_paragraph("Text after the table.")
    path = tmp_path / "sample.docx"
    document.save(path)

    blocks = read_docx(path)

    assert [block.kind for block in blocks] == ["heading", "paragraph", "table", "paragraph"]
    assert blocks[0].level == 1
    assert blocks[1].text == "Text before the table."
    assert blocks[2].rows == (("Merged heading", "Merged heading"), ("Name", "Value"))
    assert blocks[2].spans == (blocks[2].spans[0],)
    assert blocks[2].spans[0].start_row == 0
    assert blocks[2].spans[0].start_column == 0
    assert blocks[2].spans[0].end_column == 1
    assert blocks[3].text == "Text after the table."


@pytest.mark.unit
def test_html_reader_removes_index_markers_and_keeps_code_and_spans(tmp_path: Path) -> None:
    path = tmp_path / "sample.html"
    path.write_text(
        """
        <html><body>
          <h1>1. Introduction</h1>
          <p>Visible <span class="indexref">ignored index marker</span> text.</p>
          <pre>DW_OP_const1u 0x01\nDW_OP_drop</pre>
          <table><caption>Table 1. Forms</caption>
            <tr><th>Name</th><th>Value</th></tr>
            <tr><td rowspan="2">DW_FORM_addr</td><td>0x01</td></tr>
            <tr><td>continuation</td></tr>
          </table>
          <script>ignored()</script><style>.x { display: none }</style>
        </body></html>
        """,
        encoding="utf-8",
    )

    blocks = read_html(path)

    assert [block.kind for block in blocks] == ["heading", "paragraph", "code", "table"]
    assert blocks[1].text == "Visible text."
    assert blocks[2].text.strip().splitlines() == ["DW_OP_const1u 0x01", "DW_OP_drop"]
    assert blocks[3].caption == "Table 1. Forms"
    assert blocks[3].rows[1] == ("DW_FORM_addr", "0x01")
    assert blocks[3].spans[0].start_row == 1
    assert blocks[3].spans[0].end_row == 2


@pytest.mark.unit
def test_html_reader_coalesces_groff_table_header_and_continuation(tmp_path: Path) -> None:
    path = tmp_path / "groff.html"
    path.write_text(
        """
        <html><body>
          <table>
            <tr><td>lf(CW)</td></tr>
            <tr><td>Opcode Name</td></tr>
            <tr><td>Value</td></tr>
          </table>
          <p>_</p>
          <table><tr><td>DW_OP_drop</td></tr><tr><td>0x13</td></tr>
                 <tr><td>DW_OP_dup</td></tr><tr><td>0x12</td></tr></table>
        </body></html>
        """,
        encoding="utf-8",
    )

    blocks = read_html(path)

    assert len(blocks) == 1
    assert blocks[0].kind == "table"
    assert blocks[0].rows == (
        ("Opcode Name", "Value"),
        ("DW_OP_drop", "0x13"),
        ("DW_OP_dup", "0x12"),
    )
